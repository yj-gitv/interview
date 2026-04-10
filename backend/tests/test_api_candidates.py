import io

from docx import Document


def _create_position(client) -> int:
    resp = client.post(
        "/api/positions",
        json={
            "title": "产品经理",
            "jd_text": "负责用户增长",
        },
    )
    return resp.json()["id"]


def _make_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("张三", level=1)
    doc.add_paragraph("手机：13812345678")
    doc.add_paragraph("2020-2023 阿里巴巴 产品经理")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


class TestUploadCandidate:
    def test_uploads_resume_and_creates_candidate(self, client):
        pid = _create_position(client)
        docx_bytes = _make_docx_bytes()
        resp = client.post(
            f"/api/candidates/upload?position_id={pid}",
            files={
                "file": (
                    "resume.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            },
        )
        assert resp.status_code == 201
        data = resp.json()
        assert data["position_id"] == pid
        assert data["codename"] != ""

    def test_rejects_invalid_position(self, client):
        docx_bytes = _make_docx_bytes()
        resp = client.post(
            "/api/candidates/upload?position_id=999",
            files={"file": ("resume.docx", docx_bytes, "application/octet-stream")},
        )
        assert resp.status_code == 404


class TestListCandidates:
    def test_lists_by_position(self, client):
        pid = _create_position(client)
        docx_bytes = _make_docx_bytes()
        client.post(
            f"/api/candidates/upload?position_id={pid}",
            files={"file": ("r1.docx", docx_bytes, "application/octet-stream")},
        )
        resp = client.get(f"/api/candidates?position_id={pid}")
        assert resp.status_code == 200
        assert len(resp.json()) == 1


class TestGetCandidate:
    def test_gets_by_id(self, client):
        pid = _create_position(client)
        docx_bytes = _make_docx_bytes()
        create_resp = client.post(
            f"/api/candidates/upload?position_id={pid}",
            files={"file": ("r1.docx", docx_bytes, "application/octet-stream")},
        )
        cid = create_resp.json()["id"]
        resp = client.get(f"/api/candidates/{cid}")
        assert resp.status_code == 200
        assert resp.json()["id"] == cid
