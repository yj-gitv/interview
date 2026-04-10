import pytest
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.main import app
from app.database import get_db, Base

SQLALCHEMY_DATABASE_URL = "sqlite:///./test.db"

engine = create_engine(
    SQLALCHEMY_DATABASE_URL, connect_args={"check_same_thread": False}
)
TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)


@pytest.fixture(scope="function")
def db_session():
    Base.metadata.create_all(bind=engine)
    session = TestingSessionLocal()
    try:
        yield session
    finally:
        session.close()
        Base.metadata.drop_all(bind=engine)


@pytest.fixture(scope="function")
def client(db_session):
    def override_get_db():
        try:
            yield db_session
        finally:
            pass

    app.dependency_overrides[get_db] = override_get_db
    with TestClient(app) as c:
        yield c
    app.dependency_overrides.clear()


@pytest.fixture
def sample_pdf(tmp_path):
    """Create a minimal PDF for testing."""
    import fitz  # PyMuPDF

    pdf_path = tmp_path / "resume.pdf"
    doc = fitz.open()
    page = doc.new_page()
    text = (
        "张三\n"
        "手机：13812345678\n"
        "邮箱：zhangsan@example.com\n\n"
        "工作经历\n"
        "2020-2023 阿里巴巴 产品经理\n"
        "负责电商平台用户增长策略，DAU从500万提升至800万\n\n"
        "2018-2020 腾讯 产品运营\n"
        "负责微信支付商户运营\n\n"
        "教育背景\n"
        "2014-2018 北京大学 计算机科学与技术 本科\n\n"
        "技能\n"
        "数据分析、用户研究、项目管理、SQL、Python"
    )
    page.insert_text((72, 72), text, fontname="china-s", fontsize=11)
    doc.save(str(pdf_path))
    doc.close()
    return str(pdf_path)


@pytest.fixture
def sample_docx(tmp_path):
    """Create a minimal DOCX for testing."""
    from docx import Document

    docx_path = tmp_path / "resume.docx"
    doc = Document()
    doc.add_heading("李四", level=1)
    doc.add_paragraph("手机：13900001111")
    doc.add_paragraph("邮箱：lisi@example.com")
    doc.add_heading("工作经历", level=2)
    doc.add_paragraph("2021-2023 字节跳动 运营经理")
    doc.add_paragraph("负责抖音电商直播运营，GMV同比增长200%")
    doc.add_heading("教育背景", level=2)
    doc.add_paragraph("2017-2021 清华大学 工商管理 本科")
    doc.save(str(docx_path))
    return str(docx_path)
